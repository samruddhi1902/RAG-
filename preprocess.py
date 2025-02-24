import streamlit as st
# import chromadb
import os
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
import time
from langchain.docstore.document import Document as LangchainDocument
from langchain.embeddings import SentenceTransformerEmbeddings
import numpy as np
import faiss
from langchain.docstore.in_memory import InMemoryDocstore
from langchain.vectorstores import FAISS
import pinecone
from pinecone.grpc import PineconeGRPC as Pinecone
from pinecone import ServerlessSpec
import uuid
import time
def preprocess_text(files, size, overlap):
    st.write("Starting document preprocessing...")
    
    paragraphs = []
    
    for file in files:
        try:
            st.write(f"Processing file: {file.name}")
            
            if file.name.endswith(".pdf"):
                reader = PdfReader(file)
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:

                        cleaned_text = ' '.join(page_text.split())
                        paragraphs.append(cleaned_text)
                        st.write(f"Processed PDF page {i+1}, extracted {len(page_text)} characters")
                        
            elif file.name.endswith(".docx"):
                docx = DocxDocument(file)
                full_text = ""
                for para in docx.paragraphs:
                    if para.text.strip():
                        full_text += para.text.strip() + "\n\n"
                paragraphs.append(full_text)
                st.write(f"Processed DOCX file, extracted {len(full_text)} characters")
                
        except Exception as e:
            st.error(f"Error processing file {file.name}: {str(e)}")
            continue
    
    # Clean paragraphs
    paragraphs = [para for para in paragraphs if para.strip()]
    st.write(f"Total paragraphs after cleaning: {len(paragraphs)}")
    
    # Convert to Langchain Document objects
    docs = [LangchainDocument(page_content=para) for para in paragraphs]
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=size,
        chunk_overlap=overlap,
        length_function=len,
        separators=["\n\n", "\n", ". ", "? ", "! ", "; ", ": ", " ", ""],  # Order from most to least preferred split points
        is_separator_regex=False
    )
    
    # Split into chunks
    text_chunks = text_splitter.split_documents(docs)
    st.write(f"Created {len(text_chunks)} text chunks")
    
    # Debug: Print sample chunks
    if text_chunks:
        st.write("Sample chunks (first 5):")
        for i, chunk in enumerate(text_chunks[:5]):
            # Clean up the chunk display
            sample_text = chunk.page_content[:300].strip()
            st.write(f"Chunk {i+1}: {sample_text}...")
    
    return text_chunks

def preprocess_faiss(texts, embedding_model_name):
    st.write("Starting FAISS preprocessing...")  # Debug point 7
    
    embedding_model = SentenceTransformerEmbeddings(model_name=embedding_model_name)
    text_contents = [doc.page_content for doc in texts]
    
    st.write(f"Creating embeddings for {len(text_contents)} chunks...")  # Debug point 8
    embeddings = embedding_model.embed_documents(text_contents)
    embeddings = np.array(embeddings, dtype=np.float32)
    
    dimension = embeddings.shape[1]
    st.write(f"Embedding dimension: {dimension}")  # Debug point 9
    
    index = faiss.IndexFlatL2(dimension)
    index.add(embeddings)
    
    docstore = InMemoryDocstore({i: texts[i] for i in range(len(texts))})
    vector_store = FAISS(index, docstore, {}, embedding_model.embed_query)
    
    st.write("FAISS index created successfully")  # Debug point 10
    return index, docstore, vector_store

def preprocess_pinecone(text, embedding_model_name):
    import numpy as np
    from langchain.embeddings import SentenceTransformerEmbeddings
    import os
    from dotenv import load_dotenv

    load_dotenv()
    
    embedding_model = SentenceTransformerEmbeddings(model_name=embedding_model_name)

    texts = [doc.page_content for doc in text]
    embeddings = embedding_model.embed_documents(texts) 
    embeddings = np.array(embeddings)
    embeddings = embeddings.tolist()
    
    index_name = "test5"
    
    pinecone = Pinecone(
        api_key="pcsk_42Yw14_EaKdaMLiAJfWub3s2sEJYPW3jyXXjdCYkH8Mh8rD8wWJ3pS6oCCC9PGqBNuDTuf",
        environment="us-east-1")
    
    
    indexes = pinecone.list_indexes().names()
    
    if index_name in indexes:
        pinecone.delete_index(index_name)
    
    pinecone.create_index(
        name=index_name,
        dimension=len(embeddings[0]),
        metric="cosine",
        spec=ServerlessSpec(
            cloud="aws",
            region="us-east-1"
        ),
        deletion_protection="disabled"
    )

    time.sleep(5)
    
    pinecone_index = pinecone.Index(index_name)
    
    upsert_data = []
    for i in range(len(texts)):
        upsert_data.append((str(uuid.uuid4()), embeddings[i], {"text": texts[i]}))
    batch_size = 100
    for i in range(0, len(upsert_data), batch_size):
        batch = upsert_data[i:i + batch_size]
        pinecone_index.upsert(vectors=batch)
        time.sleep(0.5)
    return pinecone_index

def preprocess_vectordbs(files, embedding_model_name, size, overlap):

    if not files:
        st.warning("No files uploaded. Please upload PDF or DOCX files.")
        return None, None

    text = preprocess_text(files, size, overlap)
    if not text:
        st.warning("No valid text extracted from the files.")
        return None, None

    st.success("Preprocessing Text Complete!")

    embedding_model = SentenceTransformerEmbeddings(model_name=embedding_model_name)

    # Process FAISS
    index,docstore,vector_store = preprocess_faiss(text, embedding_model_name)
    st.success("Preprocessing Faiss Complete!")

    pinecone_index = preprocess_pinecone(text, embedding_model_name)
    st.success("Preprocessing Pinecone Complete!")

    return  index,docstore,vector_store,pinecone_index, embedding_model 
